"""
Document Converter Module
Handles conversion between different document formats (MD, PDF, DOCX, HTML)
"""

import os
import asyncio
import tempfile
from pathlib import Path
from typing import Tuple, Optional

from utils.logger import get_logger

logger = get_logger()


def _run_async(coro):
    """Run async coroutine in sync context"""
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # If loop is already running (e.g., in Qt), create a new one
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, coro)
                return future.result()
        else:
            return loop.run_until_complete(coro)
    except RuntimeError:
        return asyncio.run(coro)


class DocumentConverter:
    """Converts documents between various formats"""

    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / 'saekim_temp'
        self.temp_dir.mkdir(exist_ok=True)

    def markdown_to_pdf(self, markdown_content: str, output_path: str,
                        title: str = "Document") -> Tuple[bool, str]:
        """
        Convert Markdown to PDF using Playwright

        Args:
            markdown_content: Markdown text
            output_path: Path to save PDF
            title: Document title

        Returns:
            Tuple of (success, error_message)
        """
        try:
            # Convert markdown to HTML first
            html_content = self._markdown_to_html(markdown_content, title)

            # Use Playwright to generate PDF
            return self._generate_pdf_with_playwright(html_content, output_path)

        except Exception as e:
            error_msg = f"PDF conversion failed: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    def html_to_pdf(self, rendered_html: str, output_path: str,
                    title: str = "Document") -> Tuple[bool, str]:
        """
        Convert rendered HTML to PDF using Playwright
        This method preserves all formatting including Mermaid diagrams and KaTeX equations

        Args:
            rendered_html: Fully rendered HTML from frontend
            output_path: Path to save PDF
            title: Document title

        Returns:
            Tuple of (success, error_message)
        """
        try:
            # Wrap the rendered HTML in a complete HTML document with all dependencies
            full_html = self._create_full_html_for_pdf(rendered_html, title)

            # Use Playwright to generate PDF
            return self._generate_pdf_with_playwright(full_html, output_path)

        except Exception as e:
            error_msg = f"PDF conversion from HTML failed: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    def _create_full_html_for_pdf(self, rendered_html: str, title: str) -> str:
        """Create a complete HTML document for PDF generation with all necessary styles"""
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
    <style>
        {self._get_pdf_css()}
    </style>
</head>
<body>
    <div class="content">
        {rendered_html}
    </div>
</body>
</html>"""

    def _generate_pdf_with_playwright(self, html_content: str, output_path: str) -> Tuple[bool, str]:
        """Generate PDF from HTML using Playwright"""
        return _run_async(self._async_generate_pdf(html_content, output_path))

    async def _async_generate_pdf(self, html_content: str, output_path: str) -> Tuple[bool, str]:
        """Async implementation of PDF generation using Playwright"""
        try:
            from playwright.async_api import async_playwright
        except ImportError:
            error_msg = (
                "Playwright가 설치되지 않았습니다.\n\n"
                "다음 명령어로 설치해주세요:\n"
                "pip install playwright\n"
                "playwright install chromium"
            )
            logger.error("Playwright not installed")
            return False, error_msg

        temp_html_path = None
        try:
            # Save HTML to temp file (Playwright needs a file or URL)
            temp_html_path = self.temp_dir / f"temp_pdf_{os.getpid()}.html"
            with open(temp_html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            async with async_playwright() as p:
                # Launch browser in headless mode
                browser = await p.chromium.launch(headless=True)
                page = await browser.new_page()

                # Load the HTML file
                await page.goto(f'file:///{temp_html_path.as_posix()}', wait_until='networkidle')

                # Wait for any dynamic content (KaTeX, Mermaid) to render
                await page.wait_for_timeout(500)

                # Generate PDF with A4 format
                await page.pdf(
                    path=output_path,
                    format='A4',
                    margin={
                        'top': '2.5cm',
                        'bottom': '2.5cm',
                        'left': '2.5cm',
                        'right': '2.5cm'
                    },
                    print_background=True,
                    display_header_footer=True,
                    header_template='<div></div>',
                    footer_template='''
                        <div style="font-size: 10px; text-align: center; width: 100%; color: #666;">
                            <span class="pageNumber"></span> / <span class="totalPages"></span>
                        </div>
                    '''
                )

                await browser.close()

            logger.info(f"PDF created successfully with Playwright: {output_path}")
            return True, ""

        except Exception as e:
            error_msg = f"Playwright PDF generation failed: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

        finally:
            # Clean up temp file
            if temp_html_path and temp_html_path.exists():
                try:
                    temp_html_path.unlink()
                except:
                    pass

    def _markdown_to_html(self, markdown: str, title: str) -> str:
        """
        Convert Markdown to HTML for PDF generation
        Uses marked.js-like conversion (simplified server-side version)
        """
        # For now, use a simple conversion
        # In production, you'd use markdown library or call JS from Python
        try:
            import markdown as md_lib
            html_body = md_lib.markdown(
                markdown,
                extensions=['extra', 'codehilite', 'tables', 'fenced_code']
            )
        except ImportError:
            # Fallback: basic conversion
            html_body = self._basic_markdown_to_html(markdown)

        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
</head>
<body>
    {html_body}
</body>
</html>
"""
        return html

    def _basic_markdown_to_html(self, markdown: str) -> str:
        """
        Basic markdown to HTML conversion (fallback)
        """
        import re

        html = markdown

        # Headings
        html = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

        # Bold
        html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html)

        # Italic
        html = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html)

        # Code
        html = re.sub(r'`([^`]+)`', r'<code>\1</code>', html)

        # Links
        html = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', html)

        # Paragraphs
        html = re.sub(r'\n\n', '</p><p>', html)
        html = f'<p>{html}</p>'

        # Line breaks
        html = html.replace('\n', '<br>')

        return html

    def _get_pdf_css(self) -> str:
        """
        CSS styling for PDF output
        Enhanced to support Mermaid diagrams and KaTeX equations
        """
        return """
@page {
    size: A4;
    margin: 2.5cm;
    @top-center {
        content: "";
    }
    @bottom-center {
        content: "Page " counter(page) " of " counter(pages);
        font-size: 10pt;
        color: #666;
    }
}

body {
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}

.content {
    max-width: 100%;
}

h1 {
    font-size: 24pt;
    margin-top: 0;
    margin-bottom: 12pt;
    color: #2c3e50;
    border-bottom: 2px solid #3498db;
    padding-bottom: 6pt;
}

h2 {
    font-size: 18pt;
    margin-top: 18pt;
    margin-bottom: 9pt;
    color: #34495e;
    border-bottom: 1px solid #bdc3c7;
    padding-bottom: 3pt;
}

h3 {
    font-size: 14pt;
    margin-top: 12pt;
    margin-bottom: 6pt;
    color: #34495e;
}

p {
    margin: 6pt 0;
    text-align: justify;
}

code {
    font-family: 'Courier New', 'Consolas', monospace;
    background: #f4f4f4;
    padding: 2pt 4pt;
    border-radius: 3pt;
    font-size: 9pt;
}

pre {
    background: #f8f8f8;
    border: 1px solid #ddd;
    border-radius: 4pt;
    padding: 12pt;
    overflow-x: auto;
    font-size: 9pt;
    line-height: 1.4;
    page-break-inside: avoid;
}

pre code {
    background: transparent;
    padding: 0;
}

ul, ol {
    margin: 6pt 0;
    padding-left: 24pt;
}

li {
    margin: 3pt 0;
}

blockquote {
    margin: 12pt 0;
    padding: 6pt 12pt;
    border-left: 4px solid #3498db;
    background: #ecf0f1;
    font-style: italic;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 12pt 0;
}

table th, table td {
    border: 1px solid #bdc3c7;
    padding: 6pt 9pt;
    text-align: left;
}

table th {
    background: #34495e;
    color: white;
    font-weight: bold;
}

table tr:nth-child(even) {
    background: #ecf0f1;
}

a {
    color: #3498db;
    text-decoration: none;
}

img {
    max-width: 14cm;
    max-height: 20cm;
    height: auto;
    display: block;
    margin: 6pt auto;
    page-break-inside: avoid;
}

hr {
    border: none;
    border-top: 2px solid #bdc3c7;
    margin: 18pt 0;
}

/* Mermaid diagram styling */
.mermaid-container {
    margin: 18pt 0;
    padding: 12pt;
    background: #f9f9f9;
    border: 1px solid #e0e0e0;
    border-radius: 4pt;
    text-align: center;
    page-break-inside: avoid;
    overflow: hidden;
}

/* Mermaid as PNG images */
.mermaid-container img {
    max-width: 14cm !important;
    max-height: 20cm !important;
    width: auto !important;
    height: auto !important;
    display: block;
    margin: 0 auto;
    page-break-inside: avoid;
}

/* Fallback for SVG if not converted */
.mermaid-container svg {
    max-width: 15cm !important;
    max-height: 20cm !important;
    width: auto !important;
    height: auto !important;
    display: inline-block;
}

/* Force SVG text to use system fonts */
.mermaid-container svg text,
.mermaid-container svg tspan {
    font-family: 'Malgun Gothic', '맑은 고딕', 'Segoe UI', Arial, sans-serif !important;
    font-size: 14px !important;
}

.mermaid-error {
    color: #c00;
    background: #fee;
    padding: 12pt;
    border-radius: 4pt;
    border: 1px solid #fcc;
}

/* KaTeX math equation styling */
.katex {
    font-size: 1.1em;
}

.katex-display {
    margin: 12pt 0;
    text-align: center;
    page-break-inside: avoid;
}

.math-display {
    margin: 18pt 0;
    padding: 12pt;
    text-align: center;
    page-break-inside: avoid;
}

.math-inline {
    display: inline;
    vertical-align: baseline;
}

/* Code copy button - hide in PDF */
.code-copy-btn {
    display: none;
}

/* Code language label - hide in PDF */
.code-language-label {
    display: none;
}

/* Ensure diagrams don't get cut off */
svg {
    max-width: 15cm !important;
    max-height: 20cm !important;
    page-break-inside: avoid;
}

/* Fix SVG text rendering */
svg text,
svg tspan {
    font-family: 'Malgun Gothic', '맑은 고딕', 'Segoe UI', Arial, sans-serif !important;
}

/* Highlight.js syntax highlighting - GitHub Dark theme */
.hljs {
    display: block;
    overflow-x: auto;
    padding: 0.5em;
    color: #FFFFFF;
    background: #0d1117;
}

.hljs-doctag,
.hljs-keyword,
.hljs-meta .hljs-keyword,
.hljs-template-tag,
.hljs-template-variable,
.hljs-type,
.hljs-variable.language_ {
    color: #ff7b72;
}

.hljs-title,
.hljs-title.class_,
.hljs-title.class_.inherited__,
.hljs-title.function_ {
    color: #d2a8ff;
}

.hljs-attr,
.hljs-attribute,
.hljs-literal,
.hljs-meta,
.hljs-number,
.hljs-operator,
.hljs-selector-attr,
.hljs-selector-class,
.hljs-selector-id,
.hljs-variable {
    color: #79c0ff;
}

.hljs-meta .hljs-string,
.hljs-regexp,
.hljs-string {
    color: #a5d6ff;
}

.hljs-built_in,
.hljs-symbol {
    color: #ffa657;
}

.hljs-code,
.hljs-comment,
.hljs-formula {
    color: #8b949e;
    font-style: italic;
}

.hljs-name,
.hljs-quote,
.hljs-selector-pseudo,
.hljs-selector-tag {
    color: #7ee787;
}

.hljs-subst {
    color: #FFFFFF;
}

.hljs-section {
    color: #1f6feb;
    font-weight: bold;
}

.hljs-bullet {
    color: #f2cc60;
}

.hljs-emphasis {
    color: #FFFFFF;
    font-style: italic;
}

.hljs-strong {
    color: #FFFFFF;
    font-weight: bold;
}

.hljs-addition {
    color: #aff5b4;
    background-color: #033a16;
}

.hljs-deletion {
    color: #ffdcd7;
    background-color: #67060c;
}
"""

    def pdf_to_markdown(self, pdf_path: str) -> Tuple[bool, str, str]:
        """
        Convert PDF to Markdown (best effort)

        Args:
            pdf_path: Path to PDF file

        Returns:
            Tuple of (success, markdown_content, error_message)
        """
        try:
            # Lazy import
            import pdfplumber

            markdown_lines = []

            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text from page
                    text = page.extract_text()

                    if text:
                        markdown_lines.append(f"\n## Page {page_num}\n")
                        markdown_lines.append(text)

            markdown_content = '\n'.join(markdown_lines)
            logger.info(f"PDF converted to markdown: {pdf_path}")

            return True, markdown_content, ""

        except ImportError as e:
            error_msg = "pdfplumber is not installed"
            logger.error(error_msg)
            return False, "", error_msg
        except Exception as e:
            error_msg = f"PDF to Markdown conversion failed: {str(e)}"
            logger.error(error_msg)
            return False, "", error_msg

    def markdown_to_docx(self, markdown_content: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert Markdown to DOCX (requires python-docx)

        Args:
            markdown_content: Markdown text
            output_path: Path to save DOCX

        Returns:
            Tuple of (success, error_message)
        """
        try:
            # Lazy import
            from docx import Document

            doc = Document()

            # Basic conversion (can be enhanced)
            lines = markdown_content.split('\n')

            for line in lines:
                line = line.strip()

                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line:
                    doc.add_paragraph(line)

            doc.save(output_path)
            logger.info(f"DOCX created: {output_path}")

            return True, ""

        except ImportError as e:
            error_msg = "python-docx is not installed"
            logger.error(error_msg)
            return False, error_msg
        except Exception as e:
            error_msg = f"DOCX conversion failed: {str(e)}"
            logger.error(error_msg)
            return False, error_msg

    def markdown_to_html(self, markdown_content: str, output_path: str,
                         title: str = "Document") -> Tuple[bool, str]:
        """
        Convert Markdown to standalone HTML file

        Args:
            markdown_content: Markdown text
            output_path: Path to save HTML
            title: Document title

        Returns:
            Tuple of (success, error_message)
        """
        try:
            html_content = self._markdown_to_html(markdown_content, title)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.info(f"HTML created: {output_path}")
            return True, ""

        except Exception as e:
            error_msg = f"HTML conversion failed: {str(e)}"
            logger.error(error_msg)
            return False, error_msg
